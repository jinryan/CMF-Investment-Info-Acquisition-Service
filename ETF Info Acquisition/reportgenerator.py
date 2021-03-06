#!/usr/bin/env python

import os
cwd = os.path.dirname(os.path.realpath(__file__))
from multiprocessing import Process
import ast
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from enum import Enum
#!/usr/bin/env python

import requests
from bs4 import BeautifulSoup
import os
cwd = os.path.dirname(os.path.realpath(__file__))

def returnStockInfo(ticker, weight):
    ticker = ticker.replace(".", "-")
    industryAndSectorUrl = 'https://finance.yahoo.com/quote/' + ticker + '/profile?'
    statisticsurl = 'https://finance.yahoo.com/quote/' + ticker + '/key-statistics?p=' + ticker
    sustainabilityurl = 'https://finance.yahoo.com/quote/' + ticker + '/sustainability?p=' + ticker
    EBITAMarginURL = "https://csimarket.com/stocks/singleProfitabilityRatios.php?code=" + ticker + "&ebit"
    grossProfitMarginURL = "https://ycharts.com/companies/" + ticker + "/gross_profit_margin"
    netProfitMarginURL = "https://csimarket.com/stocks/singleProfitabilityRatios.php?code=" + ticker + "&net"


    industryAndSectorResponse = requests.get(industryAndSectorUrl)
    statisticsresponse = requests.get(statisticsurl)
    sustainabilityresponse = requests.get(sustainabilityurl)
    EBITAMarginURLResponse = requests.get(EBITAMarginURL)
    grossProfitMarginResponse = requests.get(grossProfitMarginURL)
    netProfitMarginResponse = requests.get(netProfitMarginURL)

    industryAndSectorSoup = BeautifulSoup(industryAndSectorResponse.text, "html.parser")
    etfdbsoup = BeautifulSoup(statisticsresponse.text, 'html.parser')
    sustainabilitysoup = BeautifulSoup(sustainabilityresponse.text, "html.parser")
    EBITAMarginSoup = BeautifulSoup(EBITAMarginURLResponse.text, "html.parser")
    grossProfitMarginSoup = BeautifulSoup(grossProfitMarginResponse.text, "html.parser")
    netProfitMarginSoup = BeautifulSoup(netProfitMarginResponse.text, "html.parser")

    usefulSustain = sustainabilitysoup.find("div", {"class": "smartphone_Mt(20px)", "data-reactid": "14"})
    usefulStats = sustainabilitysoup.find('section', attrs={'data-test': 'qsp-sustainability'})

    statistcs_name_box = etfdbsoup.find('section', attrs={'data-test': 'qsp-statistics'})
    sustainability_name_box = sustainabilitysoup.find('section', {'data-test': 'qsp-sustainability'})
    if statistcs_name_box == None or statisticsresponse.status_code != 200:
        return {"Status": "Data Unavailable. Please manually acquire it."}

    def findValueWithID(soup, type, id):
        s = soup.find(type, {"data-reactid": str(id)})
        if s != None:
            return s.getText().replace("\n", "").replace("\t", "").strip()
        else:
            return "N/A"

    def findIDFromFinanceSoup(id, altid=-1):
        if altid == -1:
            return findValueWithID(etfdbsoup, "td", id)
        else:
            value = findValueWithID(etfdbsoup, "td", id)
            if value == 'N/A':
                return findValueWithID(etfdbsoup, "td", altid)
            else:
                return value

    def findIDFromSustainabilitySoup(type, id):
        return findValueWithID(usefulSustain, type, id)

    def findIDFromIndustryAndSectorSoup():
        s = industryAndSectorSoup.find("p", {"data-reactid": "18"})
        try:
            sector = s.find("span", {"data-reactid": "21"}).getText().replace("\n", "").replace("\t", "").replace("&amp;", "&").strip()
            industry = s.find("span", {"data-reactid": "25"}).getText().replace("\n", "").replace("\t", "").replace("&amp;", "&").strip()
            return (sector, industry)
        except:
            return ("N/A", "N/A")

    # s = sustainabilitysoup.find("div", {"class": "smartphone_Mt(20px)", "data-reactid": "14"})
    # l = s.find("div", {"data-reactid": 20})
    attributesTuple = (
    "Ticker", "Market Capitalization", "Trailing PE", "Forward PE", "PEG Ratio (5 yr expected)", "Price/Sales (ttm)",
    "Price/Book (mrq)", "Enterprise Value/Revenue", "Enterprise Value/EBITDA", "Profit Margin", "ROA (ttm)",
    "ROE (ttm)", "Quarterly Revenue Growth (yoy)", "EBITDA", "Quarterly Earnings Growth (yoy)")

    valueDictionary = {"Ticker": ticker, "Weight": weight}
    valueDictionary["Market Capitalization"] = findIDFromFinanceSoup(19)
    valueDictionary["Trailing PE"] = findIDFromFinanceSoup(33)
    valueDictionary["Forward PE"] = findIDFromFinanceSoup(40)
    valueDictionary["PEG Ratio (5 yr expected)"] = findIDFromFinanceSoup(47)
    valueDictionary["Price/Sales (ttm)"] = findIDFromFinanceSoup(54)
    valueDictionary["Price/Book (mrq)"] = findIDFromFinanceSoup(61)
    valueDictionary["Enterprise Value/Revenue"] = findIDFromFinanceSoup(68, altid=69)
    valueDictionary["Enterprise Value/EBITDA"] = findIDFromFinanceSoup(75, altid=76)
    valueDictionary["Profit Margin"] = findIDFromFinanceSoup(112, altid=113)

    valueDictionary["ROA (ttm)"] = findIDFromFinanceSoup(131,altid=132)
    valueDictionary["ROE (ttm)"] = findIDFromFinanceSoup(138)

    valueDictionary["Quarterly Revenue Growth (yoy)"] = findIDFromFinanceSoup(164, altid=166)


    valueDictionary["EBITDA"] = findIDFromFinanceSoup(178, altid=173)

    valueDictionary["Quarterly Earnings Growth (yoy)"] = findIDFromFinanceSoup(199, altid=201)

    if sustainability_name_box == None or sustainabilityresponse.status_code != 200:
        valueDictionary["ESG Data"] = "Unavailable"
        attributesTuple += ("ESG Data",)
    else:
        attributesTuple += (
        "Total ESG Score", "Total ESG Percentile", "Environment Score", "Environment Percentile", "Social Score",
        "Social Percentile", "Governmental Score", "Governmental Percentile")
        valueDictionary["Total ESG Score"] = findIDFromSustainabilitySoup("div", 20)
        valueDictionary["Total ESG Percentile"] = findIDFromSustainabilitySoup("span", 23)

        valueDictionary["Environment Score"] = findIDFromSustainabilitySoup("div", 35)
        valueDictionary["Environment Percentile"] = findIDFromSustainabilitySoup("span", 38)

        valueDictionary["Social Score"] = findIDFromSustainabilitySoup("div", 45)
        valueDictionary["Social Percentile"] = findIDFromSustainabilitySoup("span", 48)

        valueDictionary["Governmental Score"] = findIDFromSustainabilitySoup("div", 55)
        valueDictionary["Governmental Percentile"] = findIDFromSustainabilitySoup("span", 58)

    if industryAndSectorResponse.status_code != 200:
        valueDictionary["Industry"] = "Unavailable"
        valueDictionary["Sector"] = "Unavailable"
    else:
        si = findIDFromIndustryAndSectorSoup()
        valueDictionary["Sector"] = si[0]
        valueDictionary["Industry"] = si[1]

    if EBITAMarginURLResponse.status_code != 200 or EBITAMarginSoup.find("td", {"class": "s9 zagqs sve_jedan_red dorubb"}) == None:
        valueDictionary["EBITDA Margin (Quarter)"] = "Unavailable"
    else:
        quarter = EBITAMarginSoup.find("td", {"class": "s9 zagqs sve_jedan_red dorubb"}).getText()
        quarter = quarter[quarter.find("(")+1:quarter.find(")")]
        valueDictionary["EBITDA Margin (Quarter)"] = EBITAMarginSoup.find("td", {"class": "debeligrub2 s"}).find("span").getText() + " For " + quarter

    if grossProfitMarginResponse.status_code != 200 or grossProfitMarginSoup.find("span", {"id": "pgNameVal"}) == None:
        valueDictionary["Gross Profit Margin"] = "Unavailable"
    else:
        valueDictionary["Gross Profit Margin"] = grossProfitMarginSoup.find("span", {"id": "pgNameVal"}).getText()

    if netProfitMarginResponse.status_code != 200 or netProfitMarginSoup.find("td", {"class": "s9 zagqs sve_jedan_red dorubb"}) == None:
        valueDictionary["Net Margin (Quarter)"] = "Unavailable"
    else:
        quarter = netProfitMarginSoup.find("td", {"class": "s9 zagqs sve_jedan_red dorubb"}).getText()
        quarter = quarter[quarter.find("(") + 1:quarter.find(")")]
        valueDictionary["Net Margin (Quarter)"] = netProfitMarginSoup.find("td", {
            "class": "debeligrub2 s"}).find("span").getText() + " For " + quarter


    return valueDictionary


while True:
    ticker = input("Please enter the ETF name: ")
    print("Acquring data... This may take up to 5 minutes")
    ticker = ticker.replace(".", "-")

    etfdburl = 'https://etfdb.com/etf/' + ticker + '/'
    etfurl = 'https://www.etf.com/' + ticker
    reutersurl = 'https://www.reuters.com/finance/stocks/financial-highlights/' + ticker

    etfTicker = ticker

    etfdbresponse = requests.get(etfdburl)
    etfresponse = requests.get(etfurl)
    reutersresponse = requests.get(reutersurl)


    etfdbsoup = BeautifulSoup(etfdbresponse.text, 'html.parser')
    etfsoup = BeautifulSoup(etfresponse.text, 'html.parser')
    reuterssoup = BeautifulSoup(reutersresponse.text, "html.parser")

    if etfdbresponse.status_code != 200 :
        print("We cannot search up the symbol you have entered. Please try a different one")
    elif etfresponse.status_code != 200 or reutersresponse.status_code != 200:
        print("The info will be limited.")
        break
    else:
        break


def addStNdRdTh(numberString):
    if numberString[-1] == "1":
        numberString += "st"
    elif numberString[-1] == "2":
        numberString += "nd"
    elif numberString[-1] == "3":
        numberString += "rd"
    else:
        numberString += "th"
    return numberString


# =============== BASIC INFO ===============
nameOfETF = etfdbsoup.find("h1", {"class": "data-title"}).find_all("span")[1].getText()

# =============== TOP HOLDINGS ==============
holdingsData = []

n = 4
def acquireDataForHolding(index, ticker, holdingsArray):
    global n
    nameAndTicker = ticker.getText().replace("\n", "").replace("\t", "").strip()
    print("Getting data for " + nameAndTicker)
    if "(" not in nameAndTicker:
        print(f"For {nameAndTicker}, the data is unavailable. Please manually acquire it")
        # dataAbtThatStock = {"Ticker": nameAndTicker, "Status": "Data Unavailable. Please manually acquire it."}
        n += 1
    else:
        tickerAlone = nameAndTicker[nameAndTicker.find("(")+1:nameAndTicker.find(")")]
        if tickerAlone.isalpha():
            dataAbtThatStock = returnStockInfo(tickerAlone, topHoldingsWeight[index].getText().replace("\n", "").replace("\t", "").strip())

            if type(dataAbtThatStock) is not dict:
                dataAbtThatStock = {"Ticker": nameAndTicker, "Status": "Data Unavailable. Please manually acquire it."}

            holdingsArray.append(dataAbtThatStock)

        else:
            print(f"For {nameAndTicker}, the data is unavailable. Please manually acquire it")


    print("Finished getting data for " + nameAndTicker)


print("Acquiring Data For Top 5 Holdings")
topHoldingsTicker = etfdbsoup.find_all("td", {"data-th":"Holding"})
topHoldingsWeight = etfdbsoup.find_all("td", {"data-th":"Weighting"})
procs = []

for index, ticker in enumerate(topHoldingsTicker):
    if index > n:
        break

    proc = Process(target=acquireDataForHolding(index,ticker,holdingsData, ))
    procs.append(proc)
    proc.start()

# complete the processes
for proc in procs:
    proc.join()

print("Finished Acquiring Data of Top Five Holdings")


#===================== Distribution and Expense Info =================
distributionInfo = dict()
expenseInfo = dict()
potentiallyUsefulInfo = etfdbsoup.find_all("div", {"class": "col-md-6"})
for eachPotentiallyUsefulInfo in potentiallyUsefulInfo:
    tableData = eachPotentiallyUsefulInfo.find("table", {"class": "chart base-table", "data-chart-type": "pie"})
    title = (eachPotentiallyUsefulInfo.find("h3", {"class": "h4"}))
    if tableData != None:
        # title = (eachPotentiallyUsefulInfo.find("h3", {"class": "h4"})).contents[0]
        infoArray = tableData["data-chart-series"]
        parsedArray = ast.literal_eval(infoArray)
        key = title.contents[0]
        value = []
        for eachItem in parsedArray:
            value.append(eachItem[0])
        distributionInfo[key] = value
    elif title != None:
        if title.contents[0] == "Expenses Ratio Analysis":
            expensesValues = eachPotentiallyUsefulInfo.find_all("div", {"class": "text-center"})
            expenseInfo["ETF Average"] = expensesValues[1].getText().replace("\n", "").replace("\t", "").strip()
            expenseInfo["ETF Wtd. Average"] = expensesValues[3].getText().replace("\n", "").replace("\t", "").strip()
            expenseInfo["Category Average"] = expensesValues[5].getText().replace("\n", "").replace("\t", "").strip()
            expenseRatio = eachPotentiallyUsefulInfo.find("span", {"class": "relative-metric-bubble-data"})
            expenseInfo["Expense Ratio"] = expenseRatio.getText().replace("\n", "").replace("\t", "").strip()
            # Expense Ratio Analysis



# ===================== ESG Score =================
esgScoresInfo = dict()
esgSection = etfsoup.find("div", {"id": "fundMSCIESGRatings"})
esgTitles = esgSection.find_all("label")
esgScores = esgSection.find_all("span")

for index, element in enumerate(esgTitles):
    if element == esgTitles[:-1]:
        break

    esgScoresInfo[element.getText().replace("\n", "").replace("\t", "").strip()] = esgScores[index].getText().replace("\n", "").replace("\t", "").strip()



# ====================== Performance =================
performanceInfo = dict()
esgSection = etfdbsoup.find("div", {"id": "performance"})
names = esgSection.find_all("div", {"class": "relative-metric-name"})
comparisonsWithPeer = esgSection.find_all("span", {"class", "relative-metric-rank-number"})
dataPoints = esgSection.find_all("span", {"class": "relative-metric-bubble-data"})
for index, item in enumerate(names):
    valueDict = dict()
    try:
        valueDict["Peer Ranking"] = comparisonsWithPeer[index].find("a").getText().replace("\n", "").replace("\t", "").strip()
    except:
        valueDict["Peer Ranking"] = "N/A"

    try:
        valueDict["Value"] = dataPoints[index].getText().replace("\n", "").replace("\t", "").strip()
    except:
        valueDict["Value"] = "N/A"
    performanceInfo[item.getText().replace("\n", "").replace("\t", "").strip()] = valueDict


portfolioData = etfsoup.find("div", {"id": "fundPortfolioData"})
dataPoints = portfolioData.find_all("span")




# ================== Brief Description ====================
descriptionInfo = dict()
descriptions = etfsoup.find_all("p", {"class": "pull-left mb30"})
descriptionInfo["Two Line Description"] = descriptions[0].getText().replace("\n", "").replace("\t", "").strip()
descriptionInfo["ESG Analysis"] = descriptions[2].getText().replace("\n", "").replace("\t", "").strip()
descriptionInfo["Factset Analysis"] = descriptions[1].getText().replace("\n", "").replace("\t", "").strip() # Do we need this though?

# ================== Risk Analysis ====================
riskInfo = dict()
riskSection = etfdbsoup.find("div", {"id": "technicals-collapse"})
names = riskSection.find_all("div", {"class": "relative-metric-name"})
comparisonsWithPeer = riskSection.find_all("span", {"class", "relative-metric-rank-number"})
dataPoints = riskSection.find_all("span", {"class": "relative-metric-bubble-data"})
for index, item in enumerate(names):
    valueInfo = dict()
    valueInfo["Peer Ranking"] = comparisonsWithPeer[index].getText().replace("\n", "").replace("\t", "").strip()
    valueInfo["Value"] = dataPoints[index].getText().replace("\n", "").replace("\t", "").strip()
    riskInfo[item.getText().replace("\n", "").replace("\t", "").strip()] = valueInfo



# ============= Liquidity ==============

liquidityRating = etfdbsoup.find("tr", {"id": "rc-1-liquidity"}).find("td", {"data-th": "Metric Realtime Rating"}).getText().replace("\n", "").replace("\t", "").strip()


# ============ Valuation Ratios, Dividends, Growth Rates, Financial Strengths, Profitability Ratios, Efficiency, and Management Effectiveness ============
reutersFinanceInfo = reuterssoup.find_all("table", {"cellpadding": "1", "cellspacing": "0", "class": "dataTable", "width": "100%"})
valuationRatios = dict()
valuationSection = reutersFinanceInfo[0]
valuationData = valuationSection.find_all("tr", {"class": "stripe"})
for eachDataSection in valuationData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    valuationRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance

if valuationRatios["P/E Ratio (TTM)"]["Company"] == None or valuationRatios["P/E Ratio (TTM)"]["Company"] == "--":
    valuationRatios["P/E Ratio (TTM)"]["Company"] = dataPoints[1].getText().replace("\n", "").replace("\t", "").strip()

dividendsRatios = dict()
dividendsSection = reutersFinanceInfo[1]
dividendsData = dividendsSection.find_all("tr", {"class": "stripe"})
for eachDataSection in dividendsData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    dividendsRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance



growthRatesRatios = dict()
growthRatesSection = reutersFinanceInfo[2]
growthRatesData = growthRatesSection.find_all("tr", {"class": "stripe"})
for eachDataSection in growthRatesData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    growthRatesRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance



financialStrengthsRatios = dict()
financialStrengthsSection = reutersFinanceInfo[3]
financialStrengthsData = financialStrengthsSection.find_all("tr", {"class": "stripe"})
for eachDataSection in financialStrengthsData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    financialStrengthsRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance



profitabilityRatios = dict()
profitabilitySection = reutersFinanceInfo[4]
profitabilityData = profitabilitySection.find_all("tr", {"class": "stripe"})
for eachDataSection in profitabilityData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    profitabilityRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance


efficiencyRatios = dict()
efficiencySection = reutersFinanceInfo[5]
efficiencyData = efficiencySection.find_all("tr", {"class": "stripe"})
for eachDataSection in efficiencyData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    efficiencyRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance


managementEffectivenessRatios = dict()
managementEffectivenessSection = reutersFinanceInfo[6]
managementEffectivenessData = managementEffectivenessSection.find_all("tr", {"class": "stripe"})
for eachDataSection in managementEffectivenessData:
    performance = dict()
    data = eachDataSection.find_all("td")
    performance["Company"] = data[1].getText().replace("\n", "").replace("\t", "").strip()
    performance["Industry"] = data[2].getText().replace("\n", "").replace("\t", "").strip()
    performance["Sector"] = data[3].getText().replace("\n", "").replace("\t", "").strip()
    managementEffectivenessRatios[data[0].getText().replace("\n", "").replace("\t", "").strip()] = performance



# =========== CREATE WORD DOCUMENT ============


class TableTitleStyle(Enum):
    SIDE = 1
    TOP = 2
    BOTH = 3


document = Document()
normalStyle = document.styles['Normal']
font = normalStyle.font
font.name = 'Times New Roman'
font.size = Pt(12)



# =============== Useful Functions ================
def addTitle(name):
    t1 = document.add_paragraph()
    t1.add_run(name).bold = True
    t1.style = normalStyle

def addTable(data, titleFormat=TableTitleStyle.TOP):
    # Data should be a 2 dimensional matrix
    nOfRows = len(data)
    nOfCols = len(data[0])
    table = document.add_table(rows=nOfRows, cols=nOfCols)
    table.style = 'Table Grid'

    for indexr, eachRow in enumerate(table.rows):
        for indexc, cell in enumerate(eachRow.cells):
            cell.text = data[indexr][indexc]

    if titleFormat == TableTitleStyle.TOP:
        for eachCell in table.rows[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True
    elif titleFormat == TableTitleStyle.SIDE:
        for eachCell in table.columns[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True
    else:
        for eachCell in table.rows[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True

        for eachCell in table.columns[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True

def addParagraph(paragraphText, alignment = WD_ALIGN_PARAGRAPH.LEFT):
    p = document.add_paragraph(paragraphText)
    p.style = normalStyle
    p.alignment = alignment

def newLine():
    addParagraph("")

def transpose(twoDArray):
    newArray = list()
    nC = len(twoDArray)
    nR = len(twoDArray[0])
    for i in range(nR):
        innerList = [None] * nC
        newArray.append(innerList)

    for indexi, elementi in enumerate(twoDArray):
        for indexj, elementj in enumerate(elementi):
            newArray[indexj][indexi] = elementj

    return newArray
# =========== Title ==========
titleOfDoc = document.add_heading()
ticker = etfTicker
titleOfDoc.add_run(ticker + " Investment Proposal").bold = True
titleOfDoc.style = normalStyle
titleOfDoc.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ Basic Information ==========
addTitle("Basic Information")

value = [["Name", nameOfETF], ["Ticker", ticker], ["Brief Introduction", descriptionInfo["Two Line Description"]], ["Proposed Investment Duration", "FILL THIS IN"], ["Proposed Number of Shares to Purchase", "FILL THIS IN"]]
addTable(value, titleFormat=TableTitleStyle.SIDE)
newLine()

# =================== Graph ==================
addTitle("Historical Graph and Analysis")
addParagraph("Price Trend Over Past 5 Years", alignment=WD_ALIGN_PARAGRAPH.CENTER)
addParagraph("SCREENSHOT A GRAPH HERE", alignment=WD_ALIGN_PARAGRAPH.CENTER)
newLine()
addParagraph("Historical Performance", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(performanceInfo)
peers = ["Peer Ranking"]
value = ["Value"]
for eachItem in titles:
    entries = performanceInfo[eachItem]
    peers.append(entries["Peer Ranking"])
    value.append(entries["Value"])

performanceArray = transpose([[""] + titles, peers, value])
addTable(performanceArray, titleFormat=TableTitleStyle.BOTH)
newLine()

# ==================== Behaviors ===============
addTitle("Behaviour during Bull and Bear Markets")
addParagraph("ADD SOME ANALYSIS HERE")
newLine()
addTitle("How the ETF reacted to major economic changes (e.g. trade war)")
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Risk Analysis ================
addTitle("Risk Analysis")
addParagraph("Historical Volatility", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(riskInfo)
peers = ["Peer Ranking"]
value = ["Value"]
for eachItem in titles:
    entries = riskInfo[eachItem]
    peers.append(entries["Peer Ranking"])
    value.append(entries["Value"])

riskInfoArray = transpose([["Name"] + titles, peers, value])
addTable(riskInfoArray, titleFormat=TableTitleStyle.BOTH)

addParagraph("DO SOME ANALYSIS HERE")
newLine()

# ================= Valuation ================
addTitle("Investment Rationale")
addParagraph("Part I - Valuation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(valuationRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = valuationRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

valuationArray = transpose([[""] + titles, company, industry, sector])
addTable(valuationArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Dividends ================
addParagraph("Part II - Dividends", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(dividendsRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = dividendsRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

dividendsArray = transpose([[""] + titles, company, industry, sector])
addTable(dividendsArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Profitability ================
addParagraph("Part III - Profitability", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(profitabilityRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = profitabilityRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

profitabilityArray = transpose([[""] + titles, company, industry, sector])
addTable(profitabilityArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Growth Rate Analysis  ================
addParagraph("Part IV - Growth Rate Analysis", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(growthRatesRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = growthRatesRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

growthRateArray = transpose([[""] + titles, company, industry, sector])
addTable(growthRateArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()


# ================= Management Efficiency ================
addParagraph("Part V - Management Efficiency", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(managementEffectivenessRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = managementEffectivenessRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

managementEffectivenessArray = transpose([[""] + titles, company, industry, sector])
addTable(managementEffectivenessArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================ Miscellaneous Information ===============
addParagraph("Part VI - Miscellaneous Information", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(expenseInfo)
values = ["Value"]
for eachItem in titles:
    values.append(expenseInfo[eachItem])

titles.append("Liquidity")
values.append(liquidityRating)
miscellanousArray = transpose([["Title"] + titles, values])

addTable(miscellanousArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Top Five Holdings ================
addTitle("Top Five Holdings")
titles = list(holdingsData[0])
ultimateArray = [titles]
for eachCompany in holdingsData:
    companyInfo = []
    if "Status" in list(eachCompany):
        for eachTitle in titles:
            companyInfo.append("")
    else:
        for eachTitle in titles:
            companyInfo.append(eachCompany[eachTitle])

    ultimateArray.append(companyInfo)

ultimateArray = transpose(ultimateArray)
addTable(ultimateArray, titleFormat=TableTitleStyle.SIDE)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= ESG Analysis ================
addTitle("ESG Analysis")
keys = list(esgScoresInfo)
values = list(esgScoresInfo.values())
esgArray = transpose([keys, values])
addTable(esgArray, titleFormat=TableTitleStyle.SIDE)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Summary ================
addTitle("Summary")
addParagraph("ADD SUMMARY HERE")
newLine()

print("All Done")
document.save(ticker + ' Investment Analysis.docx')